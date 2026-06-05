import json
import re
import subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palettes ────────────────────────────────────────────────────────

AO_TEXT = {
    "AO1": RGBColor(0x1F, 0x49, 0x7D),
    "AO2": RGBColor(0x37, 0x56, 0x23),
    "AO3": RGBColor(0x9C, 0x57, 0x00),
}
AO_BG = {
    "AO1": "BDD7EE",
    "AO2": "C6EFCE",
    "AO3": "FFEB9C",
}
AO_NAMES = {
    "AO1": "Knowledge & Understanding",
    "AO2": "Analysis & Evaluation",
    "AO3": "Communication",
}
GRADE_RGB = {
    1: RGBColor(0xC0, 0x00, 0x00),
    3: RGBColor(0xE2, 0x6B, 0x0A),
    5: RGBColor(0xBF, 0x8F, 0x00),
    7: RGBColor(0x2E, 0x75, 0xB6),
    9: RGBColor(0x70, 0x30, 0xA0),
}
GRADE_LABELS = {
    1: "Grade 1 — Below Standard",
    3: "Grade 3 — Developing",
    5: "Grade 5 — Solid Pass",
    7: "Grade 7 — Strong Pass",
    9: "Grade 9 — Exceptional",
}
GRADE_BG = {
    1: "FFCCCC",
    3: "FAD7A0",
    5: "FFF2CC",
    7: "DEEAF1",
    9: "E8D5F5",
}


# ── File parsing ──────────────────────────────────────────────────────────

def parse_file(filepath: str) -> str:
    p = Path(filepath)
    if p.suffix.lower() == ".pdf":
        return _parse_pdf(p)
    return _parse_docx(p)


def _parse_pdf(path: Path) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n\n".join(parts)


def _parse_docx(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── Claude subprocess helper ──────────────────────────────────────────────

def _claude(prompt: str, label: str = "") -> str:
    """Call claude -p using the Claude Code Pro plan. Retries on 529 overload."""
    import os, shutil, time

    claude_path = shutil.which("claude") or "claude"
    max_attempts = 4
    wait_seconds = [10, 20, 40]   # back-off between retries

    for attempt in range(max_attempts):
        result = subprocess.run(
            [claude_path, "-p", "--model", "opus", "--output-format", "json", prompt],
            capture_output=True,
            text=True,
            timeout=300,
            env=os.environ.copy(),
        )

        # Always try to parse stdout as JSON first — overload errors can arrive
        # with either a non-zero exit code or is_error:true, so check JSON before
        # raising on returncode.
        raw_out = result.stdout.strip()
        data = None
        if raw_out:
            try:
                data = json.loads(raw_out)
            except json.JSONDecodeError:
                pass

        # Detect overload from any path
        overload_msg = ""
        if data and data.get("is_error"):
            msg = data.get("result", data.get("error", ""))
            if "529" in str(msg) or "verload" in str(msg):
                overload_msg = str(msg)
        elif result.returncode != 0 and not data:
            raw = result.stderr or raw_out or ""
            if "529" in raw or "verload" in raw:
                overload_msg = raw

        if overload_msg:
            if attempt < max_attempts - 1:
                time.sleep(wait_seconds[attempt])
                continue
            raise RuntimeError(
                f"Claude API is overloaded ({label}). "
                "Please wait a minute and try again."
            )

        # Hard failure (non-zero exit, no parseable JSON)
        if result.returncode != 0 and not data:
            detail = (result.stderr or raw_out or "(no output)")[:600]
            raise RuntimeError(
                f"Claude CLI error{' (' + label + ')' if label else ''}: {detail}"
            )

        if not data:
            raise RuntimeError(
                f"Claude returned non-JSON ({label}): {raw_out[:400]}"
            )

        if data.get("is_error"):
            raise RuntimeError(
                f"Claude error ({label}): {data.get('result', data.get('error', ''))[:400]}"
            )

        return data.get("result", data.get("content", raw_out))

    raise RuntimeError(f"Claude did not respond after {max_attempts} attempts ({label})")


def _extract_json(text: str) -> dict:
    text = text.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    match = re.search(r"\{.*\}", text, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in Claude response")
    return json.loads(match.group())


# ── Stage 1: Document analysis ────────────────────────────────────────────

_ANALYSIS_PROMPT = """You are an expert Edexcel GCSE History examiner and marker with 15+ years experience.

Carefully read these three documents and extract precise information about the specified question.

=== QUESTION PAPER ===
{question_paper}

=== MARK SCHEME ===
{mark_scheme}

=== EXAMINER INSIGHTS / REPORT ===
{examiner_insights}

---

Extract ALL of the following information about this question: "{question_ref}"

Return ONLY valid JSON (no markdown fences, no commentary):

{{
  "question_text": "Full exact question text as printed on the paper",
  "question_type": "e.g. 16-mark essay, 4-mark explain, source evaluation",
  "context_stimulus": "Any stimulus material, source or context provided with the question (empty string if none)",
  "assessment_objectives": [
    {{
      "code": "AO1",
      "name": "Knowledge and Understanding",
      "total_marks": 8,
      "mark_descriptors": {{
        "0": "No rewardable content",
        "1-2": "Exact wording from mark scheme for 1-2 marks",
        "3-4": "Exact wording from mark scheme for 3-4 marks",
        "5-6": "Exact wording from mark scheme for 5-6 marks",
        "7-8": "Exact wording from mark scheme for 7-8 marks"
      }},
      "indicative_content": ["Key point 1 from mark scheme", "Key point 2", "Key point 3"]
    }}
  ],
  "grade_boundaries": {{
    "grade_1": {{"min": 0, "max": 5, "percentage": "31%"}},
    "grade_3": {{"min": 6, "max": 9, "percentage": "56%"}},
    "grade_5": {{"min": 10, "max": 12, "percentage": "75%"}},
    "grade_7": {{"min": 13, "max": 14, "percentage": "88%"}},
    "grade_9": {{"min": 15, "max": 16, "percentage": "94%"}}
  }},
  "grade_discriminators": {{
    "1_to_3": "What specifically lifts an answer from Grade 1 to Grade 3 (based on mark scheme language)",
    "3_to_5": "What specifically lifts an answer from Grade 3 to Grade 5",
    "5_to_7": "What specifically lifts an answer from Grade 5 to Grade 7",
    "7_to_9": "What specifically lifts an answer from Grade 7 to Grade 9"
  }},
  "examiner_report_insights": {{
    "what_worked": ["Direct quote or paraphrase from examiner report about what high-scoring answers did"],
    "common_errors": ["Direct quote or paraphrase about what lower-scoring answers did wrong"],
    "specific_advice": "Any specific examiner advice for this question"
  }},
  "total_marks": 16
}}

Be precise. Use the exact language from the mark scheme descriptors, not paraphrases."""


def analyze_documents(
    question_paper: str,
    mark_scheme: str,
    examiner_insights: str,
    question_ref: str,
) -> dict:
    prompt = _ANALYSIS_PROMPT.format(
        question_paper=question_paper[:9000],
        mark_scheme=mark_scheme[:7000],
        examiner_insights=examiner_insights[:4000],
        question_ref=question_ref,
    )
    raw = _claude(prompt, "analysis")
    return _extract_json(raw)


# ── Stage 2: Exemplar generation ──────────────────────────────────────────

_GENERATION_PROMPT = """You are a highly experienced Edexcel GCSE History examiner generating authentic student exemplar answers.

You have already analysed the question and mark scheme. Use this analysis to generate grade-accurate exemplars:

=== MARK SCHEME ANALYSIS ===
{analysis}

---

Generate authentic exemplar answers for these grades: {grades_list}

CRITICAL RULES — read carefully before generating:

1. EACH GRADE NEEDS ITS OWN AUTHENTIC VOICE. Do NOT write a Grade 9 answer and water it down.
   Each grade represents a fundamentally different student: different vocabulary, different analytical depth,
   different evidence selection, different sentence structures.

2. GRADE ANCHORS:
   - Grade 1: Basic recall only. Simple sentences. May name an event but gives no analysis.
     Uses vague language ("lots of", "many", "important"). 1-2 pieces of evidence maximum, used superficially.
   - Grade 3: Shows some understanding. Attempts analysis but stays descriptive ("this caused X because...").
     Evidence is relevant but lacks precision. Argument is linear with no evaluation.
   - Grade 5: Clear analytical engagement. Evidence is specific (dates, names, statistics).
     Begins to evaluate or acknowledge alternative views. Structured argument. Vocabulary shows understanding.
   - Grade 7: Sophisticated analysis. Evidence deployed strategically to support the argument.
     Genuine evaluation — compares significance of factors or weighs evidence.
     Nuanced argument with counterargument acknowledged.
   - Grade 9: Evaluative mastery. Conceptual framing of the answer. Evidence is precise and selective.
     Explicitly evaluates the relative weight of factors. Handles complexity and historiography where relevant.
     Elegant, precise communication. Every sentence advances the argument.

3. GROUND EVERY SEGMENT in the exact mark descriptor language from the analysis.
   When writing Grade 5 AO2, it should demonstrably meet the 5-6 mark AO2 descriptor — not exceed it.

4. USE the examiner report insights: avoid the common errors noted, replicate what the high-scoring answers did.

5. SEGMENTS should be paragraph-length chunks (not single sentences, not multiple paragraphs).
   Assign the PRIMARY assessment objective to each segment.

Return ONLY valid JSON (no markdown fences):

{{
  "question_text": "Full question text",
  "question_type": "Question type",
  "assessment_objectives": [
    {{"code": "AO1", "name": "Knowledge and Understanding", "total_marks": 8}},
    {{"code": "AO2", "name": "Analysis and Evaluation", "total_marks": 6}},
    {{"code": "AO3", "name": "Communication", "total_marks": 2}}
  ],
  "exemplars": [
    {{
      "grade": 5,
      "total_marks": 11,
      "mark_breakdown": {{"AO1": 5, "AO2": 4, "AO3": 2}},
      "mark_justification": {{
        "AO1": "5/8 — References relevant knowledge with some specific detail, meeting the 3-4 band descriptor: [quote exact descriptor]. Falls short of 5-6 because...",
        "AO2": "4/6 — Analysis is clear but evaluation is limited, meeting the 3-4 band: [quote exact descriptor].",
        "AO3": "2/2 — Spelling, punctuation and grammar are accurate throughout."
      }},
      "grade_characteristics_met": [
        "Specific characteristic from the grade anchor that this answer demonstrates"
      ],
      "grade_characteristics_missed": [
        "Specific characteristic from the next grade up that this answer lacks"
      ],
      "segments": [
        {{
          "text": "Full paragraph text of the student answer",
          "primary_ao": "AO1",
          "annotation": "Short label: what this segment demonstrates (e.g. Precise factual knowledge with dates)",
          "examiner_note": "2-3 sentences. Which exact mark descriptor band does this meet and why? What has the student done that earns marks here? Reference the mark scheme language directly.",
          "gap_to_full_marks": "1-2 precise sentences explaining what is MISSING or UNDERDEVELOPED in this segment that prevents it reaching the top of its band. Be explicit: what specific knowledge, analytical move, or evaluative step would be needed? If this segment already achieves maximum available marks for its objective, write: 'Achieves maximum marks for this objective.'"
        }}
      ]
    }}
  ]
}}"""


def generate_exemplars(analysis: dict, grades: list) -> dict:
    prompt = _GENERATION_PROMPT.format(
        analysis=json.dumps(analysis, indent=2)[:8000],
        grades_list=", ".join(f"Grade {g}" for g in grades),
    )
    raw = _claude(prompt, "generation")
    return _extract_json(raw)


# ── Word document helpers ─────────────────────────────────────────────────

def _run(para, text, *, bold=False, italic=False, size=None, rgb=None):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if rgb:
        r.font.color.rgb = rgb
    return r


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_col_width(table, col_idx: int, width_cm: float):
    for row in table.rows:
        cell = row.cells[col_idx]
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement("w:tcW")
        tcW.set(qn("w:w"), str(int(width_cm * 567)))
        tcW.set(qn("w:type"), "dxa")
        tcPr.append(tcW)


# ── Word document builder ─────────────────────────────────────────────────

def build_word_doc(data: dict, output_path: Path):
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(1.8)
        sec.bottom_margin = Cm(1.8)
        sec.left_margin   = Cm(2.2)
        sec.right_margin  = Cm(1.8)

    _build_cover(doc, data)
    for exemplar in data.get("exemplars", []):
        doc.add_page_break()
        _build_exemplar(doc, exemplar)

    doc.save(output_path)


def _build_cover(doc: Document, data: dict):
    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(h, "GCSE History — Edexcel", bold=True, size=18, rgb=RGBColor(0x1F, 0x49, 0x7D))

    sub = doc.add_heading("", level=2)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(sub, "Examiner Exemplar Answers", bold=True, size=14, rgb=RGBColor(0x1F, 0x49, 0x7D))

    doc.add_paragraph()

    qp = doc.add_paragraph()
    _run(qp, "QUESTION:  ", bold=True, size=11, rgb=RGBColor(0x1F, 0x49, 0x7D))
    _run(qp, data.get("question_text", ""), size=11)

    ctx = data.get("context_stimulus", "")
    if ctx:
        cp = doc.add_paragraph()
        cp.paragraph_format.left_indent = Cm(0.8)
        _run(cp, "Context: ", bold=True, size=10, rgb=RGBColor(0x55, 0x55, 0x55))
        _run(cp, ctx, italic=True, size=10, rgb=RGBColor(0x55, 0x55, 0x55))

    total = sum(ao.get("total_marks", 0) for ao in data.get("assessment_objectives", []))
    info = doc.add_paragraph()
    _run(info, f"Type: {data.get('question_type', '')}   |   Total marks: {total}",
         italic=True, size=10, rgb=RGBColor(0x55, 0x55, 0x55))

    doc.add_paragraph()

    ao_hdr = doc.add_paragraph()
    _run(ao_hdr, "Assessment Objectives", bold=True, size=11)
    for ao in data.get("assessment_objectives", []):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.8)
        code = ao.get("code", "")
        _run(p, f"{code}  —  {ao.get('name', '')}  ({ao.get('total_marks', 0)} marks)",
             size=10, rgb=AO_TEXT.get(code, RGBColor(0, 0, 0)))

    doc.add_paragraph()

    legend_hdr = doc.add_paragraph()
    _run(legend_hdr, "Annotation Colour Key", bold=True, size=10)

    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, (code, name) in enumerate(AO_NAMES.items()):
        cell = tbl.cell(0, i)
        _set_cell_bg(cell, AO_BG[code])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, f"{code}\n", bold=True, size=10, rgb=AO_TEXT[code])
        _run(p, name, size=9, rgb=AO_TEXT[code])


def _build_exemplar(doc: Document, exemplar: dict):
    grade   = exemplar.get("grade", 0)
    g_rgb   = GRADE_RGB.get(grade, RGBColor(0, 0, 0))
    g_bg    = GRADE_BG.get(grade, "FFFFFF")

    h = doc.add_heading("", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(h, GRADE_LABELS.get(grade, f"Grade {grade}"), bold=True, size=20, rgb=g_rgb)

    mp = doc.add_paragraph()
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total = exemplar.get("total_marks", 0)
    bd    = exemplar.get("mark_breakdown", {})
    bd_str = "   |   ".join(f"{k}: {v}" for k, v in bd.items())
    _run(mp, f"Total: {total} marks     {bd_str}", bold=True, size=12, rgb=g_rgb)

    doc.add_paragraph()

    for item in exemplar.get("grade_characteristics_met", []):
        p = doc.add_paragraph()
        _run(p, "✔  ", bold=True, size=10, rgb=RGBColor(0x37, 0x56, 0x23))
        _run(p, item, size=10, rgb=RGBColor(0x37, 0x56, 0x23))

    for item in exemplar.get("grade_characteristics_missed", []):
        p = doc.add_paragraph()
        _run(p, "▲  To reach next grade: ", bold=True, size=10, rgb=RGBColor(0xC0, 0x00, 0x00))
        _run(p, item, size=10, rgb=RGBColor(0xC0, 0x00, 0x00))

    # ── Continuous answer (clean prose, no annotations) ──────────────────────
    cont_hdr = doc.add_paragraph()
    _run(cont_hdr, "FULL STUDENT ANSWER", bold=True, size=9,
         rgb=RGBColor(0x55, 0x55, 0x55))

    cont_tbl = doc.add_table(rows=1, cols=1)
    cont_tbl.style = "Table Grid"
    _set_col_width(cont_tbl, 0, 16.5)

    cont_cell = cont_tbl.cell(0, 0)
    _set_cell_bg(cont_cell, "FAFAFA")

    segments = exemplar.get("segments", [])
    for i, seg in enumerate(segments):
        p = cont_cell.paragraphs[0] if i == 0 else cont_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0) if i == 0 else Pt(10)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.3)
        p.paragraph_format.right_indent = Cm(0.3)
        _run(p, seg.get("text", ""), size=11)

    doc.add_paragraph()

    # ── Separator ─────────────────────────────────────────────────────────────
    sep_p = doc.add_paragraph()
    sep_p.paragraph_format.space_before = Pt(4)
    sep_p.paragraph_format.space_after  = Pt(4)
    _run(sep_p, "─" * 110, size=7, rgb=RGBColor(0xBB, 0xBB, 0xBB))

    col_hdr = doc.add_paragraph()
    _run(col_hdr, "EXAMINER ANNOTATIONS — SECTION BY SECTION", bold=True, size=9,
         rgb=RGBColor(0x55, 0x55, 0x55))

    for seg in exemplar.get("segments", []):
        ao    = seg.get("primary_ao", "AO1")
        ao_rgb = AO_TEXT.get(ao, RGBColor(0, 0, 0))
        ao_bg  = AO_BG.get(ao, "FFFFFF")

        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        _set_col_width(tbl, 0, 11.0)
        _set_col_width(tbl, 1, 5.5)

        ac = tbl.cell(0, 0)
        _set_cell_bg(ac, "FFFFFF")
        ap = ac.paragraphs[0]
        ap.paragraph_format.space_before = Pt(4)
        ap.paragraph_format.space_after  = Pt(4)
        _run(ap, seg.get("text", ""), size=11)

        nc = tbl.cell(0, 1)
        _set_cell_bg(nc, ao_bg)
        np_ = nc.paragraphs[0]
        _run(np_, f"[{ao}]  ", bold=True, size=9, rgb=ao_rgb)
        _run(np_, seg.get("annotation", ""), bold=True, size=9, rgb=ao_rgb)

        ep = nc.add_paragraph()
        ep.paragraph_format.space_before = Pt(3)
        _run(ep, seg.get("examiner_note", ""), italic=True, size=8,
             rgb=RGBColor(0x40, 0x40, 0x40))

        gap = seg.get("gap_to_full_marks", "").strip()
        if gap:
            achieves_max = gap.lower().startswith("achieves maximum")

            sep_p = nc.add_paragraph()
            sep_p.paragraph_format.space_before = Pt(5)
            _run(sep_p, "─" * 28, size=7, rgb=RGBColor(0xBB, 0xBB, 0xBB))

            if achieves_max:
                lbl = nc.add_paragraph()
                lbl.paragraph_format.space_before = Pt(2)
                _run(lbl, "✔  FULL MARKS ACHIEVED", bold=True, size=8,
                     rgb=RGBColor(0x37, 0x56, 0x23))
                gp = nc.add_paragraph()
                _run(gp, gap, italic=True, size=8, rgb=RGBColor(0x37, 0x56, 0x23))
            else:
                lbl = nc.add_paragraph()
                lbl.paragraph_format.space_before = Pt(2)
                _run(lbl, "▲  WHY NOT FULL MARKS:", bold=True, size=8,
                     rgb=RGBColor(0xC0, 0x00, 0x00))
                gp = nc.add_paragraph()
                gp.paragraph_format.space_before = Pt(2)
                _run(gp, gap, size=8, rgb=RGBColor(0xC0, 0x00, 0x00))

        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Mark justification table
    doc.add_paragraph()
    mj_h = doc.add_paragraph()
    _run(mj_h, "Examiner Mark Justification", bold=True, size=11, rgb=g_rgb)

    mj = exemplar.get("mark_justification", {})
    mj_tbl = doc.add_table(rows=len(mj) + 1, cols=2)
    mj_tbl.style = "Table Grid"
    _set_col_width(mj_tbl, 0, 2.5)
    _set_col_width(mj_tbl, 1, 14.0)

    hr = mj_tbl.rows[0]
    for cell in hr.cells:
        _set_cell_bg(cell, "1F497D")
    _run(hr.cells[0].paragraphs[0], "Objective",    bold=True, size=10, rgb=RGBColor(0xFF, 0xFF, 0xFF))
    _run(hr.cells[1].paragraphs[0], "Justification", bold=True, size=10, rgb=RGBColor(0xFF, 0xFF, 0xFF))

    for i, (code, justification) in enumerate(mj.items()):
        row = mj_tbl.rows[i + 1]
        _set_cell_bg(row.cells[0], AO_BG.get(code, "FFFFFF"))
        _set_cell_bg(row.cells[1], "FFFFFF")
        _run(row.cells[0].paragraphs[0], code, bold=True, size=10, rgb=AO_TEXT.get(code, RGBColor(0, 0, 0)))
        _run(row.cells[1].paragraphs[0], justification, size=10)
